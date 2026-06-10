from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Pt
from latex2mathml.converter import convert
from lxml import etree


ROOT = Path(r"E:\AUDI_COLOR\RGB")
OUT_PATH = ROOT / "test_white_mix_数学推导.docx"
MML2OMML_XSL = Path(r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL")


def qn(tag: str) -> str:
    prefix, tag_root = tag.split(":")
    uri = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }[prefix]
    return f"{{{uri}}}{tag_root}"


def set_run_fonts(run, latin="Times New Roman", east_asia="Microsoft YaHei", size=11, bold=False):
    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_asia)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_fonts(run, size=18, bold=True)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_fonts(run, size=14 if level == 1 else 12, bold=True)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_fonts(run, size=11)


class EquationWriter:
    def __init__(self, xsl_path: Path):
        xslt = etree.parse(str(xsl_path))
        self.transform = etree.XSLT(xslt)

    def latex_to_omml(self, latex: str):
        mathml = convert(latex)
        mathml_tree = etree.fromstring(mathml.encode("utf-8"))
        omml_tree = self.transform(mathml_tree)
        return omml_tree.getroot()

    def add_equation(self, doc: Document, latex: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        omml = self.latex_to_omml(latex)
        p._p.append(parse_xml(etree.tostring(omml)))


def add_equation_block(doc, writer, *equations):
    for eq in equations:
        writer.add_equation(doc, eq)


def main():
    doc = Document()
    writer = EquationWriter(MML2OMML_XSL)

    add_title(doc, "test_white_mix.m 数学推导")
    add_body(doc, "本文对应当前 RGB 文件夹下 test_white_mix.m、xyY_PWM.m 与 RGB_to_xyY_32000.m 的实际实现流程。")
    add_body(doc, "下面的公式使用 Word 原生数学公式对象写入，分式、矩阵、上下标均为标准排版。")

    add_heading(doc, "1. 已知量")
    add_equation_block(
        doc,
        writer,
        r"R=(x_r,y_r,Y_r)=(0.69,0.31,1.40)",
        r"G=(x_g,y_g,Y_g)=(0.17,0.72,2.20)",
        r"B=(x_b,y_b,Y_b)=(0.15,0.03,0.50)",
        r"W=(x_W,y_W)=(0.3090,0.3176)",
        r"\mathrm{MaxPWM}=32000",
        r"T=(x_T,y_T,Y_T)",
        r"T_{xy}=(x_T,y_T)",
    )

    add_heading(doc, "2. 白点射线与色域交点")
    add_body(doc, "从白点 W 出发，经过目标点 T_{xy} 作射线，并与色域三角形边界求交。")
    add_equation_block(
        doc,
        writer,
        r"L(t)=W+t\left(T_{xy}-W\right),\quad t\ge 1",
        r"E_i(s)=A_i+s(B_i-A_i),\quad 0\le s\le 1",
        r"W+t\left(T_{xy}-W\right)=A_i+s(B_i-A_i)",
        r"P=(x_P,y_P)",
    )
    add_body(doc, "在全部满足条件的交点中，取满足 t\ge 1 且 t 最小的那个交点作为有效交点。")

    add_heading(doc, "3. 共线参数")
    add_equation_block(
        doc,
        writer,
        r"k_1=\left|P-T_{xy}\right|",
        r"k_2=\left|T_{xy}-W\right|",
        r"k_3=\left|P-W\right|",
        r"\lambda=\frac{k_2}{k_3}",
        r"1-\lambda=\frac{k_1}{k_3}",
        r"T_{xy}=(1-\lambda)W+\lambda P",
        r"x_T=(1-\lambda)x_W+\lambda x_P",
        r"y_T=(1-\lambda)y_W+\lambda y_P",
    )

    add_heading(doc, "4. 目标亮度分解")
    add_body(doc, "根据当前代码采用的亮度推导，将目标亮度 Y_T 分解成白点分量亮度与交点分量亮度。")
    add_equation_block(
        doc,
        writer,
        r"Y_W=Y_T\frac{(1-\lambda)y_W}{y_T}",
        r"Y_P=Y_T\frac{\lambda y_P}{y_T}",
        r"Y_W+Y_P=Y_T",
        r"W'=(x_W,y_W,Y_W)",
        r"P'=(x_P,y_P,Y_P)",
    )

    add_heading(doc, "5. xyY\\_PWM.m 的第一步：求三基色系数 D")
    add_equation_block(
        doc,
        writer,
        r"z_r=1-x_r-y_r,\qquad z_g=1-x_g-y_g,\qquad z_b=1-x_b-y_b",
        r"z_C=1-x_C-y_C",
        r"M=\begin{bmatrix}"
        r"\frac{Y_r}{y_r}x_r & \frac{Y_g}{y_g}x_g & \frac{Y_b}{y_b}x_b \\"
        r"\frac{Y_r}{y_r}y_r & \frac{Y_g}{y_g}y_g & \frac{Y_b}{y_b}y_b \\"
        r"\frac{Y_r}{y_r}z_r & \frac{Y_g}{y_g}z_g & \frac{Y_b}{y_b}z_b"
        r"\end{bmatrix}",
        r"D=\begin{bmatrix}D_r\\D_g\\D_b\end{bmatrix}"
        r"=M^{-1}\begin{bmatrix}x_C\\y_C\\z_C\end{bmatrix}",
    )
    add_body(doc, "在当前实现里，C 可以取白点分量 W'，也可以取交点分量 P'。")

    add_heading(doc, "6. xyY\\_PWM.m 的第二步：由 D 求最终 PWM")
    add_equation_block(
        doc,
        writer,
        r"L_C=Y_rD_r+Y_gD_g+Y_bD_b",
        r"R_C=\frac{Y_C}{L_C}\,\mathrm{MaxPWM}\,D_r",
        r"G_C=\frac{Y_C}{L_C}\,\mathrm{MaxPWM}\,D_g",
        r"B_C=\frac{Y_C}{L_C}\,\mathrm{MaxPWM}\,D_b",
        r"PWM_W=\begin{bmatrix}R_W\\G_W\\B_W\end{bmatrix},\qquad"
        r"PWM_P=\begin{bmatrix}R_P\\G_P\\B_P\end{bmatrix}",
    )

    add_heading(doc, "7. 两束光的 PWM 合成")
    add_equation_block(
        doc,
        writer,
        r"PWM_{\mathrm{mix}}=PWM_W+PWM_P",
        r"R_{\mathrm{mix}}=R_W+R_P,\qquad G_{\mathrm{mix}}=G_W+G_P,\qquad B_{\mathrm{mix}}=B_W+B_P",
    )

    add_heading(doc, "8. 直接送入 RGB_to_xyY_32000.m 的回算输入")
    add_body(doc, "当前版本不再把 PWM 换算成等效 0–255 RGB，而是直接使用 0–32000 标尺的 PWM 作为回算输入。")
    add_equation_block(
        doc,
        writer,
        r"BackInput=\begin{bmatrix}R_{\mathrm{mix}}\\G_{\mathrm{mix}}\\B_{\mathrm{mix}}\end{bmatrix}",
    )

    add_heading(doc, "9. RGB_to_xyY_32000.m 的回算验证")
    add_body(doc, "记直接调用 RGB_to_xyY_32000.m 后得到的验证点为 V=(x_V,y_V,Y_V)。")
    add_equation_block(
        doc,
        writer,
        r"S=\frac{Y_r}{y_r}R_{\mathrm{mix}}+\frac{Y_g}{y_g}G_{\mathrm{mix}}+\frac{Y_b}{y_b}B_{\mathrm{mix}}",
        r"x_V=\frac{\frac{Y_r}{y_r}x_rR_{\mathrm{mix}}+\frac{Y_g}{y_g}x_gG_{\mathrm{mix}}+\frac{Y_b}{y_b}x_bB_{\mathrm{mix}}}{S}",
        r"y_V=\frac{Y_rR_{\mathrm{mix}}+Y_gG_{\mathrm{mix}}+Y_bB_{\mathrm{mix}}}{S}",
        r"Y_V=\frac{Y_rR_{\mathrm{mix}}+Y_gG_{\mathrm{mix}}+Y_bB_{\mathrm{mix}}}{32000}",
    )

    add_heading(doc, "10. 误差验证")
    add_body(doc, "以下的 u' 与 v' 对应 CIE 1976 u'v' 坐标。")
    add_equation_block(
        doc,
        writer,
        r"u_T=\frac{4x_T}{-2x_T+12y_T+3},\qquad v_T=\frac{9y_T}{-2x_T+12y_T+3}",
        r"u_V=\frac{4x_V}{-2x_V+12y_V+3},\qquad v_V=\frac{9y_V}{-2x_V+12y_V+3}",
        r"\Delta uv=\sqrt{(u_V-u_T)^2+(v_V-v_T)^2}",
        r"\Delta Y=\left|Y_V-Y_T\right|",
        r"\delta_Y=\frac{\Delta Y}{Y_T}",
    )

    add_heading(doc, "11. 过程总式")
    add_body(doc, "把 test_white_mix.m 当前实现压缩成一条链式表达，可写成：")
    add_equation_block(
        doc,
        writer,
        r"T\rightarrow (W,P,\lambda)\rightarrow (Y_W,Y_P)\rightarrow (W',P')"
        r"\rightarrow (PWM_W,PWM_P)\rightarrow PWM_{\mathrm{mix}}"
        r"\rightarrow (x_V,y_V,Y_V)",
    )

    try:
        doc.save(str(OUT_PATH))
        print(f"Saved: {OUT_PATH}")
    except PermissionError:
        fallback_path = ROOT / "test_white_mix_数学推导_32000版.docx"
        doc.save(str(fallback_path))
        print(f"Saved: {fallback_path}")


if __name__ == "__main__":
    main()
